import os
import re
import json
import numpy as np
import networkx as nx
from functools import lru_cache
from typing import Generator

from langchain_community.document_loaders import PDFPlumberLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.llms import Ollama
from langchain.prompts import PromptTemplate
from langchain.retrievers import BM25Retriever, EnsembleRetriever
from sentence_transformers import CrossEncoder
from rag_engine_graph_optimized import build_graph_rag_fast, GRAPH_CONFIG

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

# ══════════════════════════════════════════════════════════════
# CẤU HÌNH TẬP TRUNG
# ══════════════════════════════════════════════════════════════

CONFIG = {
    "embedding_model": "sentence-transformers/paraphrase-multilingual-mpnet-base-v2",
    "embedding_device": "cpu",
    "chunk_size": 1000,
    "chunk_overlap": 100,
    "retriever_k": 5,
    "llm_model": "qwen2.5:7b",
    "llm_temperature": 0.3,
    "ollama_host": os.environ.get("OLLAMA_HOST", "http://localhost:11434"),
    "use_rerank": False,
    "use_self_rag": False,
    "self_rag_max_iter": 2,
    "selected_file": "All",
}

_VIET_CHARS = "àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ"

# ══════════════════════════════════════════════════════════════
# 1. EMBEDDING MODEL & LLM — SINGLETON CÓ CACHE
# ══════════════════════════════════════════════════════════════

_embedder = None
_cross_encoder = None
_llm_instance = None
_llm_config_key = None
# Alias để không phải đổi tên ở các chỗ khác
build_graph_rag = build_graph_rag_fast

def get_embedder() -> HuggingFaceEmbeddings:
    global _embedder
    if _embedder is None:
        _embedder = HuggingFaceEmbeddings(
            model_name=CONFIG["embedding_model"],
            model_kwargs={"device": CONFIG["embedding_device"]},
            encode_kwargs={"normalize_embeddings": True},
        )
    return _embedder


def get_cross_encoder() -> CrossEncoder:
    global _cross_encoder
    if _cross_encoder is None:
        _cross_encoder = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
    return _cross_encoder


def _get_llm() -> Ollama:
    """Chỉ tạo lại khi config thay đổi (model / host / temperature)."""
    global _llm_instance, _llm_config_key
    current_key = (CONFIG["llm_model"], CONFIG["ollama_host"], CONFIG["llm_temperature"])
    if _llm_instance is None or _llm_config_key != current_key:
        _llm_instance = Ollama(
            model=CONFIG["llm_model"],
            base_url=CONFIG["ollama_host"],
            temperature=CONFIG["llm_temperature"],
        )
        _llm_config_key = current_key
    return _llm_instance


# ══════════════════════════════════════════════════════════════
# 2. XỬ LÝ TÀI LIỆU
# ══════════════════════════════════════════════════════════════

def _clean_text(text: str) -> str:
    """
    Làm sạch text trước khi đưa vào chunk/context.
    - Xóa ký tự control không in được
    - Chuẩn hóa nhiều newline liên tiếp
    - Xóa khoảng trắng thừa
    - Xóa dòng chỉ có số trang
    """
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'(?m)^\s{0,2}\d{1,4}\s*$', '', text)
    return text.strip()


def process_pdf(file_path: str, embedder, chunk_size: int = None, chunk_overlap: int = None) -> tuple:
    """Đọc PDF → clean → chunk. Trả về: (chunks, num_pages, num_chunks)"""
    loader = PDFPlumberLoader(file_path)
    docs = loader.load()

    for doc in docs:
        doc.page_content = _clean_text(doc.page_content)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size or CONFIG["chunk_size"],
        chunk_overlap=chunk_overlap or CONFIG["chunk_overlap"],
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
    )
    chunks = splitter.split_documents(docs)
    chunks = [c for c in chunks if len(c.page_content.strip()) >= 50]
    return chunks, len(docs), len(chunks)


def process_docx(file_path: str, embedder, chunk_size: int = None, chunk_overlap: int = None) -> tuple:
    """Đọc DOCX → clean → chunk. Trả về: (chunks, num_pages, num_chunks)"""
    from langchain_community.document_loaders import Docx2txtLoader
    import docx

    CHARS_PER_PAGE = 1000

    loader = Docx2txtLoader(file_path)
    docs = loader.load()

    for doc in docs:
        doc.page_content = _clean_text(doc.page_content)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size or CONFIG["chunk_size"],
        chunk_overlap=chunk_overlap or CONFIG["chunk_overlap"],
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
    )
    chunks = splitter.split_documents(docs)
    chunks = [c for c in chunks if len(c.page_content.strip()) >= 50]

    try:
        doc_obj = docx.Document(file_path)
        total_chars = sum(len(p.text.strip()) for p in doc_obj.paragraphs if p.text.strip())
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_chars += len(cell.text.strip())
        num_pages = max(1, round(total_chars / CHARS_PER_PAGE))
    except Exception:
        total_chars = sum(len(d.page_content) for d in docs)
        num_pages = max(1, round(total_chars / CHARS_PER_PAGE))

    return chunks, num_pages, len(chunks)


def build_hybrid_retriever(chunks, embedder):
    """
    Xây Hybrid Retriever (FAISS + BM25).
    Trọng số 0.3 BM25 / 0.7 FAISS — semantic search quan trọng hơn.
    """
    vector_store = FAISS.from_documents(chunks, embedder)
    faiss_retriever = vector_store.as_retriever(
        search_type="mmr",
        search_kwargs={"k": CONFIG["retriever_k"], "fetch_k": CONFIG["retriever_k"] * 3},
    )
    bm25_retriever = BM25Retriever.from_documents(chunks)
    bm25_retriever.k = CONFIG["retriever_k"]

    hybrid_retriever = EnsembleRetriever(
        retrievers=[bm25_retriever, faiss_retriever],
        weights=[0.3, 0.7],
    )
    return hybrid_retriever


# ══════════════════════════════════════════════════════════════
# 3. GRAPH RAG
# ══════════════════════════════════════════════════════════════

def _extract_entities_and_relations_llm(text: str, llm) -> list[dict]:
    """
    Dùng LLM để trích xuất entities và relationships.

    CẢI TIẾN:
    - Tăng text từ 500 lên 800 chars để không bỏ sót thông tin quan trọng
    - Phân loại relationship type rõ ràng (CAUSES, CONTAINS, REQUIRES, ...)
    - Thêm rule chống entity generic cho tiếng Việt
    - Thêm minimum 3 chars per entity
    """
    prompt = f"""Extract key entities and their relationships from the Vietnamese/English text below.
Return ONLY a JSON array. No explanation, no markdown, no preamble.

Format: [{{"e1": "entity1", "rel": "relationship_type:verb_phrase", "e2": "entity2"}}]

Relationship types (use one prefix):
- CAUSES: (e1 gây ra / dẫn đến e2)
- CONTAINS: (e1 bao gồm / chứa e2)
- REQUIRES: (e1 cần / phụ thuộc e2)
- DEFINES: (e1 là định nghĩa / mô tả e2)
- OPPOSES: (e1 mâu thuẫn / trái với e2)
- IMPLEMENTS: (e1 thực hiện / áp dụng e2)

Rules:
- e1 and e2 must be SPECIFIC named concepts, tools, methods, organizations, regulations, or processes
- FORBIDDEN generic words: "hệ thống", "phương pháp", "dữ liệu", "thông tin", "quy trình",
  "system", "method", "data", "information", "process", "thing", "item", "value"
- Minimum 3 characters per entity name
- Extract 3 to 8 relationships maximum
- Keep entities in their ORIGINAL language from the text (do not translate)

Text:
{text[:800]}

JSON array only:"""

    try:
        raw = llm.invoke(prompt).strip()
        raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
        match = re.search(r'\[.*\]', raw, re.DOTALL)
        if match:
            raw = match.group()
        relations = json.loads(raw)
        valid = [
            r for r in relations
            if isinstance(r, dict)
            and all(k in r for k in ("e1", "rel", "e2"))
            and r["e1"].strip() and r["e2"].strip()
            and len(r["e1"]) > 2 and len(r["e2"]) > 2
        ]
        return valid[:8]
    except Exception:
        # Fallback: extract acronym + Title Case
        entities = list(set(
            re.findall(r'\b[A-Z]{2,}\b', text) +
            re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b', text)
        ))[:6]
        relations = []
        for i in range(len(entities)):
            for j in range(i + 1, len(entities)):
                relations.append({"e1": entities[i], "rel": "DEFINES:related_to", "e2": entities[j]})
        return relations[:6]


def _build_graph_rag_legacy(chunks: list) -> nx.Graph:
    """Xây Knowledge Graph từ chunks bằng LLM-based entity extraction."""
    llm = _get_llm()
    embedder = get_embedder()
    G = nx.Graph()

    for idx, chunk in enumerate(chunks):
        text = chunk.page_content
        page = chunk.metadata.get("page", "?")
        source = chunk.metadata.get("source", "unknown")
        chunk_data = {
            "chunk_idx": idx,
            "page": page,
            "source": source,
            "content": text,
        }

        relations = _extract_entities_and_relations_llm(text, llm)

        for rel in relations:
            e1 = rel["e1"].strip()
            e2 = rel["e2"].strip()
            relationship = rel["rel"].strip()

            if G.has_node(e1):
                G.nodes[e1]["chunks"].append(chunk_data)
            else:
                G.add_node(e1, chunks=[chunk_data], label=e1)

            if G.has_node(e2):
                G.nodes[e2]["chunks"].append(chunk_data)
            else:
                G.add_node(e2, chunks=[chunk_data], label=e2)

            if G.has_edge(e1, e2):
                G[e1][e2]["weight"] += 1
                if relationship not in G[e1][e2]["rels"]:
                    G[e1][e2]["rels"].append(relationship)
            else:
                G.add_edge(e1, e2, weight=1, rels=[relationship])

    # Tính embedding cho từng node (batch)
    node_list = list(G.nodes())
    if node_list:
        try:
            node_embeddings = embedder.embed_documents(node_list)
            for node, emb in zip(node_list, node_embeddings):
                G.nodes[node]["embedding"] = emb
        except Exception:
            pass

    return G


def _score_chunk(content: str, keywords: list) -> float:
    if not keywords or not content:
        return 0.0
    content_lower = content.lower()
    word_count = max(len(content_lower.split()), 1)
    hits = sum(1 for kw in keywords if kw in content_lower)
    if hits == 0:
        return 0.0
    coverage  = hits / len(keywords)
    hits_norm = min(hits / len(keywords), 1.0)
    density   = min(hits / (word_count / 50), 1.0)
    return 0.5 * coverage + 0.3 * hits_norm + 0.2 * density


def _graph_retrieve(question: str, graph: nx.Graph, top_k: int = 6) -> list[dict]:
    """
    Graph RAG Retrieval — 3 tầng:
    Tầng 1: Semantic Node Matching (embedding similarity)
    Tầng 2: Graph Traversal (neighbor expansion)
    Tầng 3: Chunk Deduplication + Ranking
    """
    MIN_SIM       = 0.25
    MAX_SEED      = 5
    MAX_NEIGHBORS = 3

    node_list = list(graph.nodes())
    if not node_list:
        return []

    nodes_with_emb = [
        (n, graph.nodes[n]["embedding"])
        for n in node_list
        if "embedding" in graph.nodes[n]
    ]

    seed_nodes: dict = {}

    if nodes_with_emb:
        embedder = get_embedder()
        q_emb = np.array(embedder.embed_query(question))
        q_norm = np.linalg.norm(q_emb)

        for node, node_emb in nodes_with_emb:
            n_emb = np.array(node_emb)
            n_norm = np.linalg.norm(n_emb)
            if q_norm == 0 or n_norm == 0:
                continue
            sim = float(np.dot(q_emb, n_emb) / (q_norm * n_norm))
            if sim >= MIN_SIM:
                seed_nodes[node] = sim

        seed_nodes = dict(
            sorted(seed_nodes.items(), key=lambda x: x[1], reverse=True)[:MAX_SEED]
        )
    else:
        keywords = [
            w.lower() for w in re.findall(r'[\w\u00C0-\u024F\u1E00-\u1EFF]+', question)
            if len(w) >= 4
        ]
        for node in node_list:
            nl = node.lower()
            if any(kw in nl or nl in kw for kw in keywords):
                seed_nodes[node] = 0.7

    if not seed_nodes:
        keywords = [
            w.lower() for w in re.findall(r'[\w\u00C0-\u024F\u1E00-\u1EFF]+', question)
            if len(w) >= 3
        ]
        all_chunks: dict = {}
        for node in node_list:
            for ci in graph.nodes[node].get("chunks", []):
                idx = ci["chunk_idx"]
                if idx not in all_chunks:
                    s = _score_chunk(ci["content"], keywords)
                    all_chunks[idx] = (s, ci)
        ranked_fb = sorted(all_chunks.values(), key=lambda x: x[0], reverse=True)
        return [ci for _, ci in ranked_fb[:top_k] if _ > 0]

    neighbor_nodes: dict = {}
    for node, sim in seed_nodes.items():
        nbrs = list(graph.neighbors(node))
        if not nbrs:
            continue
        nbrs_sorted = sorted(
            nbrs,
            key=lambda n: graph[node][n].get("weight", 0),
            reverse=True,
        )
        max_w = graph[node][nbrs_sorted[0]].get("weight", 1) if nbrs_sorted else 1
        for nb in nbrs_sorted[:MAX_NEIGHBORS]:
            if nb in seed_nodes:
                continue
            edge_w = graph[node][nb].get("weight", 1)
            nb_score = sim * 0.6 * (edge_w / max_w)
            if nb not in neighbor_nodes or neighbor_nodes[nb] < nb_score:
                neighbor_nodes[nb] = nb_score

    chunk_scores: dict = {}

    for node, sim in seed_nodes.items():
        if not graph.has_node(node):
            continue
        for ci in graph.nodes[node].get("chunks", []):
            idx = ci["chunk_idx"]
            if idx not in chunk_scores or chunk_scores[idx][0] < sim:
                chunk_scores[idx] = (sim, ci)

    for node, nb_score in neighbor_nodes.items():
        if not graph.has_node(node):
            continue
        for ci in graph.nodes[node].get("chunks", []):
            idx = ci["chunk_idx"]
            if idx not in chunk_scores or chunk_scores[idx][0] < nb_score:
                chunk_scores[idx] = (nb_score, ci)

    ranked = sorted(chunk_scores.values(), key=lambda x: x[0], reverse=True)
    return [ci for _, ci in ranked[:top_k]]


# ══════════════════════════════════════════════════════════════
# 4. HELPERS NỘI BỘ
# ══════════════════════════════════════════════════════════════

def _detect_language(text: str) -> str:
    return "vi" if any(c in text.lower() for c in _VIET_CHARS) else "en"


def _build_prompt(language: str) -> PromptTemplate:
    """
    PROMPT CHÍNH — Classic RAG.

    - Thêm chain-of-thought 2 bước: đọc/phân tích → trả lời
    - Yêu cầu trích dẫn số trang cụ thể để tăng độ tin cậy
    - Xử lý rõ trường hợp nhiều chunk mâu thuẫn nhau
    - Định dạng output theo từng loại câu hỏi (điểm đơn / nhiều điểm / so sánh)
    - Cấm bịa đặt số liệu, tên người, ngày tháng
    - KHÔNG yêu cầu độ dài cứng nhắc — thay bằng hướng dẫn theo ngữ cảnh
    """
    if language == "vi":
        template = """Bạn là chuyên gia phân tích tài liệu nội bộ. Nhiệm vụ của bạn là trả lời câu hỏi CHỈ dựa trên các đoạn trích được cung cấp.

═══ BƯỚC 1 — ĐỌC VÀ PHÂN TÍCH NGỮ CẢNH ═══
Trước khi trả lời, hãy xác định nội tâm:
• Đoạn nào chứa thông tin trực tiếp trả lời câu hỏi?
• Đoạn nào cung cấp bối cảnh bổ sung?
• Có mâu thuẫn giữa các đoạn không? Nếu có, ưu tiên đoạn chi tiết hơn hoặc xuất hiện sau.

═══ BƯỚC 2 — TRẢ LỜI ═══
NGUYÊN TẮC BẮT BUỘC:
1. CHỈ dùng thông tin trong [NGỮ CẢNH]. Tuyệt đối không thêm kiến thức bên ngoài.
2. Khi trích dẫn thông tin quan trọng, ghi rõ nguồn: (Trang X) hoặc (Đoạn Y).
3. Nếu nhiều đoạn đề cập cùng một vấn đề → tổng hợp, không lặp lại.
4. Nếu ngữ cảnh có một phần thông tin → trả lời phần có, ghi rõ: "Tài liệu không đề cập đến [phần thiếu]."
5. Nếu ngữ cảnh hoàn toàn không liên quan → trả lời: "Tài liệu không đề cập đến vấn đề này."
6. Với câu hỏi follow-up → kết hợp [HỘI THOẠI TRƯỚC] và [NGỮ CẢNH] để giữ ngữ mạch liên tục.
7. KHÔNG bịa đặt số liệu, tên người, ngày tháng, hoặc bất kỳ dữ kiện nào không có trong tài liệu.
8. KHÔNG bắt đầu bằng "Dựa vào ngữ cảnh..." hay lặp lại câu hỏi.

ĐỊNH DẠNG TRẢ LỜI:
• Câu hỏi có 1 điểm chính → đoạn văn trả lời trực tiếp, súc tích.
• Câu hỏi có nhiều điểm → dùng gạch đầu dòng hoặc đánh số, mỗi điểm 1-2 câu.
• Câu hỏi so sánh → 2 đoạn đối chiếu hoặc bảng rõ ràng.

[HỘI THOẠI TRƯỚC]
{chat_history}

[NGỮ CẢNH]
{context}

[CÂU HỎI]
{question}

[TRẢ LỜI — Phân tích từ ngữ cảnh, trích dẫn trang khi có thể]"""

    else:
        template = """You are an internal document analysis expert. Your task is to answer the question using ONLY the provided excerpts.

═══ STEP 1 — READ AND ANALYZE THE CONTEXT ═══
Before answering, mentally identify:
• Which passage directly answers the question?
• Which passages provide supporting context?
• Are there contradictions between passages? If so, prioritize the more detailed or later-appearing one.

═══ STEP 2 — ANSWER ═══
MANDATORY RULES:
1. ONLY use information in [CONTEXT]. Never add outside knowledge.
2. When citing important information, reference the source: (Page X) or (Passage Y).
3. If multiple passages cover the same topic → synthesize, do not repeat.
4. If context has partial info → answer what you can, clearly state: "The document does not mention [missing part]."
5. If context is completely unrelated → respond: "The document does not mention this topic."
6. For follow-up questions → combine [HISTORY] and [CONTEXT] for a coherent, continuous answer.
7. NEVER fabricate figures, names, dates, or any fact not present in the document.
8. DO NOT start with "Based on the context..." or repeat the question.

RESPONSE FORMAT:
• Single-point question → direct prose answer, concise.
• Multi-point question → bullet points or numbered list, 1-2 sentences each.
• Comparison question → two contrasting paragraphs or a clear table.

[HISTORY]
{chat_history}

[CONTEXT]
{context}

[QUESTION]
{question}

[ANSWER — Analyze from context, cite page numbers where possible]"""

    return PromptTemplate(template=template, input_variables=["context", "question", "chat_history"])


def _build_graph_prompt(language: str) -> PromptTemplate:
    """
    PROMPT GRAPH RAG.

    CẢI TIẾN SO VỚI PHIÊN BẢN CŨ:
    - Giải thích rõ cho model biết đặc điểm của Graph RAG retrieval
      (chunks được chọn qua mạng quan hệ thực thể, không phải chỉ similarity)
    - Hướng dẫn 4 bước phân tích tường minh: xác định → tìm → truy vết → tổng hợp
    - Yêu cầu ghi nguồn (trang) khi trích dẫn
    - Xử lý rõ trường hợp mâu thuẫn giữa các đoạn
    """
    if language == "vi":
        template = """Bạn là chuyên gia phân tích tài liệu sử dụng Graph RAG. Các đoạn dưới đây được chọn dựa trên MẠNG LƯỚI QUAN HỆ giữa các thực thể — nghĩa là chúng được liên kết với nhau qua các khái niệm, không chỉ đơn thuần gần về nghĩa.

═══ ĐẶC ĐIỂM CỦA NGỮ CẢNH NÀY ═══
Bạn nhận được {num_chunks} đoạn văn được kết nối qua đồ thị tri thức. Các đoạn này:
• Có thể đến từ nhiều phần khác nhau trong tài liệu
• Được chọn vì chứa các thực thể LIÊN QUAN đến câu hỏi qua mối quan hệ trực tiếp hoặc gián tiếp
• Cùng nhau tạo thành bức tranh đầy đủ hơn bất kỳ đoạn nào riêng lẻ

═══ NHIỆM VỤ 4 BƯỚC ═══
1. XÁC ĐỊNH các thực thể/khái niệm chính trong câu hỏi.
2. TÌM các thực thể đó trong các đoạn văn được cung cấp.
3. TRUY VẾT mối quan hệ: A dẫn đến B, A phụ thuộc vào C, B mâu thuẫn với D.
4. TỔNG HỢP thông tin từ nhiều đoạn — ưu tiên phân tích SỰ KẾT NỐI hơn liệt kê đơn thuần.

NGUYÊN TẮC BẮT BUỘC:
1. CHỈ dùng thông tin trong [NGỮ CẢNH]. Không thêm kiến thức bên ngoài.
2. Khi trích dẫn thông tin quan trọng, ghi rõ nguồn: (Trang X) hoặc (Đoạn Y).
3. Nếu các đoạn mâu thuẫn → ghi nhận cả hai quan điểm và giải thích sự khác biệt.
4. Nếu không có thông tin liên quan → trả lời: "Tài liệu không đề cập đến vấn đề này."
5. KHÔNG bắt đầu bằng "Dựa vào ngữ cảnh..." hay lặp lại câu hỏi.
6. KHÔNG bịa đặt số liệu, tên người, ngày tháng không có trong tài liệu.

[HỘI THOẠI TRƯỚC]
{chat_history}

[NGỮ CẢNH — {num_chunks} đoạn liên kết qua đồ thị quan hệ]
{context}

[CÂU HỎI]
{question}

[TRẢ LỜI — Phân tích quan hệ thực thể, tổng hợp từ nhiều đoạn, ghi nguồn khi cần]"""

    else:
        template = """You are a document analysis expert using Graph RAG. The passages below were selected based on the RELATIONSHIP NETWORK between entities — meaning they are connected through shared concepts, not just semantic similarity alone.

═══ CHARACTERISTICS OF THIS CONTEXT ═══
You receive {num_chunks} passages connected via a knowledge graph. These passages:
• May come from different sections of the document
• Were selected because they contain entities RELATED to the question through direct or indirect relationships
• Together form a more complete picture than any single passage alone

═══ 4-STEP TASK ═══
1. IDENTIFY the key entities/concepts in the question.
2. LOCATE those entities in the provided passages.
3. TRACE the relationships: A leads to B, A depends on C, B contradicts D.
4. SYNTHESIZE information across passages — prioritize analyzing CONNECTIONS over simple listing.

MANDATORY RULES:
1. ONLY use information in [CONTEXT]. No outside knowledge.
2. When citing important information, reference the source: (Page X) or (Passage Y).
3. If passages contradict each other → acknowledge both views and explain the difference.
4. If no relevant info found → respond: "The document does not mention this topic."
5. DO NOT start with "Based on the context..." or repeat the question.
6. NEVER fabricate figures, names, or dates not present in the document.

[HISTORY]
{chat_history}

[CONTEXT — {num_chunks} passages linked via relationship graph]
{context}

[QUESTION]
{question}

[ANSWER — Analyze entity relationships, synthesize across passages, cite sources when needed]"""

    return PromptTemplate(
        template=template,
        input_variables=["context", "question", "chat_history", "num_chunks"],
    )


def _format_chat_history(chat_history: list, max_turns: int = 4) -> str:
    """
    Format lịch sử hội thoại để đưa vào prompt.

    CẢI TIẾN SO VỚI PHIÊN BẢN CŨ:
    - Tăng max_turns từ 3 lên 4 để có thêm ngữ cảnh hội thoại
    - Tăng answer_truncate từ 200 lên 350 chars — tránh mất ngữ cảnh câu trả lời dài
    - Dùng role label [Người dùng]/[Trợ lý] rõ ràng hơn User:/Assistant:
      giúp model phân biệt role tốt hơn trong tiếng Việt
    """
    if not chat_history:
        return "Không có lịch sử hội thoại."

    history_text = []
    recent = chat_history[-max_turns:]
    answer_truncate = 350

    for item in recent:
        q = item.get("question", "").strip()
        a = item.get("answer", "").strip()
        if not q:
            continue

        # Skip compare-mode answers (chứa marker đặc biệt)
        if "**Classic RAG:**" in a or "**Graph RAG:**" in a:
            continue

        # Truncate answer dài nhưng giữ nhiều hơn để bảo toàn ngữ cảnh
        if len(a) > answer_truncate:
            a = a[:answer_truncate].rsplit(" ", 1)[0] + "..."

        history_text.append(f"[Người dùng]: {q}")
        history_text.append(f"[Trợ lý]: {a}")

    return "\n".join(history_text) if history_text else "Không có lịch sử hội thoại."


def _build_context(source_docs) -> str:
    """
    Build context có đánh số đoạn văn rõ ràng.
    Giúp LLM dễ tham chiếu khi trả lời (Đoạn 1, Trang X...).
    """
    parts = []
    for i, doc in enumerate(source_docs, 1):
        page = doc.metadata.get("page", "?")
        if isinstance(page, int):
            page = page + 1
        source = os.path.basename(doc.metadata.get("source", ""))
        header = f"[Đoạn {i} — Trang {page}" + (f" | {source}" if source else "") + "]"
        parts.append(f"{header}\n{doc.page_content.strip()}")
    return "\n\n---\n\n".join(parts)


def rerank_documents(question: str, docs: list, top_k: int = 5):
    if not docs:
        return docs
    cross_encoder = get_cross_encoder()
    pairs = [(question, doc.page_content) for doc in docs]
    scores = cross_encoder.predict(pairs)
    for doc, score in zip(docs, scores):
        doc.metadata["rerank_score"] = float(score)
    docs = sorted(docs, key=lambda x: x.metadata["rerank_score"], reverse=True)
    return docs[:top_k]


# ══════════════════════════════════════════════════════════════
# 5. HỎI ĐÁP — CLASSIC RAG (streaming)
# ══════════════════════════════════════════════════════════════

def ask_question_stream_with_sources(
    question: str,
    retriever,
    chat_history: list = None,
    file_name: str = None,
) -> Generator[str, None, None]:

    lang = _detect_language(question)
    prompt = _build_prompt(lang)
    history_text = _format_chat_history(chat_history or [])

    # 1. Retrieve candidates
    retrieved_docs = retriever.invoke(question)[:20]

    # 2. Filter theo selected_file
    selected_file = CONFIG.get("selected_file", "All")
    if selected_file != "All":
        retrieved_docs = [
            d for d in retrieved_docs
            if d.metadata.get("source") == selected_file
        ]

    # 3. Self-RAG path
    if CONFIG.get("use_self_rag", False):
        answer, score, context, docs = self_rag_pipeline(question, retriever, chat_history)
        yield answer
        sources = [
            {
                "index": i + 1,
                "page": doc.metadata.get("page", "?"),
                "source": doc.metadata.get("source", "unknown"),
                "content": doc.page_content,
            }
            for i, doc in enumerate(docs)
        ]
        yield "@@SOURCES@@" + json.dumps(sources, ensure_ascii=False)
        yield "@@CONFIDENCE@@" + str(score)
        return

    # 4. Rerank hoặc lấy top-k
    if CONFIG.get("use_rerank", False):
        source_docs = rerank_documents(question, retrieved_docs, top_k=CONFIG["retriever_k"])
    else:
        source_docs = retrieved_docs[:CONFIG["retriever_k"]]

    if not source_docs:
        no_info = "Không tìm thấy thông tin liên quan trong tài liệu." if lang == "vi" else "No relevant information found in the document."
        yield no_info
        yield "@@SOURCES@@" + json.dumps([], ensure_ascii=False)
        return

    context = _build_context(source_docs)

    filled_prompt = prompt.format(
        context=context,
        question=question,
        chat_history=history_text,
    )

    llm = _get_llm()
    try:
        for token in llm.stream(filled_prompt):
            if token:
                yield token
    except Exception as e:
        yield f"\n\n[Lỗi khi gọi model: {e}]"
        return

    sources = [
        {
            "index": i + 1,
            "page": (doc.metadata.get("page", "?") + 1)
                    if isinstance(doc.metadata.get("page"), int)
                    else doc.metadata.get("page", "?"),
            "source": file_name or os.path.basename(doc.metadata.get("source", "unknown")),
            "content": doc.page_content,
        }
        for i, doc in enumerate(source_docs)
    ]
    yield "@@SOURCES@@" + json.dumps(sources, ensure_ascii=False)


# ══════════════════════════════════════════════════════════════
# 6. HỎI ĐÁP — GRAPH RAG (non-streaming)
# ══════════════════════════════════════════════════════════════

def ask_graph_rag(
    question: str,
    graph: nx.Graph,
    chat_history: list = None,
) -> tuple[str, list]:
    """Trả lời câu hỏi bằng Graph RAG. Trả về: (answer_string, sources_list)"""
    lang = _detect_language(question)
    prompt = _build_graph_prompt(lang)
    history_text = _format_chat_history(chat_history or [])

    graph_chunks = _graph_retrieve(question, graph, top_k=CONFIG["retriever_k"])

    if not graph_chunks:
        no_info = "Graph RAG không tìm thấy thông tin liên quan trong đồ thị." if lang == "vi" else "Graph RAG found no relevant information."
        return no_info, []

    context = "\n\n---\n\n".join(
        f"[Đoạn {i+1} — Trang {c.get('page', '?')}]\n{c['content'].strip()}"
        for i, c in enumerate(graph_chunks)
    )
    filled_prompt = prompt.format(
        context=context,
        question=question,
        chat_history=history_text,
        num_chunks=len(graph_chunks),
    )

    llm = _get_llm()
    answer = llm.invoke(filled_prompt)

    sources = [
        {
            "index": i + 1,
            "page": c.get("page", "?"),
            "source": os.path.basename(str(c.get("source", "unknown"))),
            "content": c["content"],
        }
        for i, c in enumerate(graph_chunks)
    ]

    return answer, sources


# ══════════════════════════════════════════════════════════════
# 7. SO SÁNH CLASSIC RAG vs GRAPH RAG (generator)
# ══════════════════════════════════════════════════════════════

def ask_compare_rag(
    question: str,
    retriever,
    graph: nx.Graph,
    chat_history: list = None,
) -> Generator[str, None, None]:
    """Stream song song Classic RAG và Graph RAG để so sánh."""
    lang = _detect_language(question)
    history_text = _format_chat_history(chat_history or [])
    llm = _get_llm()

    # ── Phần 1: Classic RAG ──────────────────────────────────
    yield "@@CLASSIC_START@@"

    classic_docs = retriever.invoke(question)[:20]
    selected_file = CONFIG.get("selected_file", "All")
    if selected_file != "All":
        classic_docs = [d for d in classic_docs if d.metadata.get("source") == selected_file]

    if CONFIG.get("use_rerank", False):
        source_docs = rerank_documents(question, classic_docs, top_k=CONFIG["retriever_k"])
    else:
        source_docs = classic_docs[:CONFIG["retriever_k"]]

    context_classic = _build_context(source_docs)
    classic_prompt = _build_prompt(lang)
    filled_classic = classic_prompt.format(
        context=context_classic,
        question=question,
        chat_history=history_text,
    )

    for token in llm.stream(filled_classic):
        if token:
            yield token

    classic_sources = [
        {
            "index": i + 1,
            "page": (doc.metadata.get("page", "?") + 1)
                    if isinstance(doc.metadata.get("page"), int)
                    else doc.metadata.get("page", "?"),
            "source": os.path.basename(str(doc.metadata.get("source", "unknown"))),
            "content": doc.page_content,
        }
        for i, doc in enumerate(source_docs)
    ]
    yield "@@CLASSIC_SOURCES@@" + json.dumps(classic_sources, ensure_ascii=False)

    # ── Phần 2: Graph RAG ────────────────────────────────────
    yield "@@GRAPH_START@@"

    graph_top_k = CONFIG["retriever_k"] + 3
    graph_chunks = _graph_retrieve(question, graph, top_k=graph_top_k)

    if graph_chunks:
        context_graph = "\n\n---\n\n".join(
            f"[Đoạn {i+1} — Trang {c.get('page', '?')}]\n{c['content'].strip()}"
            for i, c in enumerate(graph_chunks)
        )
        graph_prompt = _build_graph_prompt(lang)
        filled_graph = graph_prompt.format(
            context=context_graph,
            question=question,
            chat_history=history_text,
            num_chunks=len(graph_chunks),
        )
        for token in llm.stream(filled_graph):
            if token:
                yield token
    else:
        no_info = "Graph RAG không tìm thấy thông tin liên quan trong đồ thị." if lang == "vi" else "Graph RAG found no relevant information."
        yield no_info
        graph_chunks = []

    graph_sources = [
        {
            "index": i + 1,
            "page": c.get("page", "?"),
            "source": os.path.basename(str(c.get("source", "unknown"))),
            "content": c["content"],
        }
        for i, c in enumerate(graph_chunks)
    ]
    yield "@@GRAPH_SOURCES@@" + json.dumps(graph_sources, ensure_ascii=False)
    yield "@@COMPARE_DONE@@"


# ══════════════════════════════════════════════════════════════
# 8. SELF-RAG
# ══════════════════════════════════════════════════════════════

def rewrite_query(question: str, chat_history: str = "") -> str:
    """
    Viết lại câu hỏi để tăng khả năng tìm kiếm trong hệ thống RAG.

    CẢI TIẾN SO VỚI PHIÊN BẢN CŨ:
    - Cung cấp 5 chiến lược rewrite tường minh thay vì chỉ nói "more specific"
    - Thêm hướng dẫn giữ nguyên ngôn ngữ gốc (tiếng Việt vẫn là tiếng Việt)
    - Thêm rule: nếu câu hỏi đã đủ cụ thể → giữ nguyên, không rewrite thừa
    - Thêm hướng dẫn expand abbreviations và thêm loại thực thể nếu bị ngầm hiểu
    """
    llm = _get_llm()
    history_part = (
        f"\nLịch sử hội thoại gần nhất:\n{chat_history}\n"
        if chat_history.strip() and chat_history != "Không có lịch sử hội thoại."
        else ""
    )
    prompt = f"""Bạn là trình tối ưu hóa truy vấn cho hệ thống tìm kiếm tài liệu (RAG).
Nhiệm vụ: Viết lại câu hỏi để tăng tối đa khả năng tìm đúng đoạn văn trong cơ sở dữ liệu tài liệu.

Chiến lược viết lại (áp dụng cái phù hợp nhất):
1. MỞ RỘNG từ viết tắt hoặc ký hiệu ngắn gọn
2. THÊM loại thực thể nếu bị ngầm hiểu (vd: "điều khoản" → "điều khoản hợp đồng lao động")
3. TÁCH câu hỏi phức hợp thành câu hỏi con có thể tìm kiếm nhất
4. THAY THẾ đại từ (nó, họ, điều này...) bằng tên cụ thể từ lịch sử hội thoại
5. GIỮ NGUYÊN ngôn ngữ gốc (tiếng Việt vẫn là tiếng Việt, không dịch sang tiếng Anh)

Nếu câu hỏi đã đủ cụ thể và rõ ràng → xuất nguyên văn, không thay đổi.
Chỉ xuất câu hỏi đã viết lại. Không giải thích, không thêm tiền tố, không dùng dấu ngoặc kép.{history_part}

Câu hỏi gốc: {question}
Câu hỏi đã viết lại:"""

    result = llm.invoke(prompt).strip()
    # Nếu LLM trả lời nhiều dòng, chỉ lấy dòng đầu
    return result.split("\n")[0].strip() or question


def evaluate_answer(question: str, answer: str, context: str) -> dict:
    """
    Đánh giá chất lượng câu trả lời cho vòng lặp Self-RAG.

    CẢI TIẾN SO VỚI PHIÊN BẢN CŨ:
    - Thêm rubric chấm điểm 5 mức rõ ràng (0.9-1.0, 0.7-0.8, 0.5-0.6, 0.3-0.4, 0.0-0.2)
    - Thêm field "missing" để biết thông tin còn thiếu → giúp rewrite tốt hơn
    - Tăng ctx_short từ 600 lên 800 chars để context đủ cho model đánh giá
    - Tăng ans_short từ 300 lên 400 chars
    """
    llm = _get_llm()
    ctx_short = context[:800] if len(context) > 800 else context
    ans_short = answer[:400] if len(answer) > 400 else answer

    prompt = f"""Đánh giá chất lượng câu trả lời dựa trên câu hỏi và ngữ cảnh.

Rubric chấm điểm:
- 0.9-1.0: Câu trả lời đầy đủ, chính xác, trả lời trực tiếp câu hỏi, mọi thông tin đều có trong ngữ cảnh
- 0.7-0.8: Câu trả lời phần lớn đúng, thiếu sót nhỏ, được hỗ trợ tốt bởi ngữ cảnh
- 0.5-0.6: Câu trả lời chỉ trả lời một phần, có thông tin không được ngữ cảnh xác nhận, hoặc bỏ sót điểm chính
- 0.3-0.4: Câu trả lời mơ hồ, phần lớn không liên quan, hoặc mâu thuẫn với ngữ cảnh
- 0.0-0.2: Câu trả lời sai, bịa đặt, hoặc hoàn toàn lạc đề

Chỉ trả về JSON: {{"score": <0.0-1.0>, "missing": "<thông tin còn thiếu trong câu trả lời, nếu có>", "reason": "<1 câu giải thích>"}}

Câu hỏi: {question}
Ngữ cảnh (trích đoạn): {ctx_short}
Câu trả lời: {ans_short}

JSON:"""

    result = llm.invoke(prompt).strip()
    try:
        match = re.search(r'\{.*\}', result, re.DOTALL)
        if match:
            return json.loads(match.group())
        return json.loads(result)
    except Exception:
        return {"score": 0.5, "missing": "", "reason": "parse_error"}


def self_rag_pipeline(question, retriever, chat_history=None):
    """
    Self-RAG: tự đánh giá và cải thiện câu trả lời qua nhiều vòng lặp.

    Vòng 1: dùng câu hỏi gốc
    Vòng 2+: rewrite query dựa trên thông tin "missing" từ evaluate_answer
    """
    history_text = _format_chat_history(chat_history or [])
    best_answer, best_score, best_context, best_docs = "", 0.0, "", []

    for iteration in range(CONFIG["self_rag_max_iter"]):
        # Vòng 1: dùng câu hỏi gốc; vòng 2+: rewrite dựa trên missing info
        if iteration == 0:
            search_query = question
        else:
            # Truyền thêm thông tin "missing" vào lịch sử để rewrite thông minh hơn
            missing_hint = f"\n(Lưu ý: câu trả lời trước còn thiếu: {best_eval.get('missing', '')})" if best_eval.get("missing") else ""
            search_query = rewrite_query(question + missing_hint, history_text)

        docs = retriever.invoke(search_query)[:10]

        if CONFIG.get("use_rerank", False):
            docs = rerank_documents(search_query, docs, top_k=CONFIG["retriever_k"])
        else:
            docs = docs[:CONFIG["retriever_k"]]

        context = _build_context(docs)
        prompt = _build_prompt(_detect_language(question))
        filled = prompt.format(context=context, question=question, chat_history=history_text)

        llm = _get_llm()
        answer = llm.invoke(filled)
        best_eval = evaluate_answer(question, answer, context)
        score = best_eval.get("score", 0.0)

        if score > best_score:
            best_score = score
            best_answer = answer
            best_context = context
            best_docs = docs

        if score > 0.8:
            break

    return best_answer, best_score, best_context, best_docs